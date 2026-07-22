"""DOCX document extractor."""

import logging
import re
from io import BytesIO

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.run import Run

from tools.document import Document, ExtractorResult
from tools.extractor_base import BaseExtractor
from tools.helpers import render_markdown_table
from tools.image_assets import is_http_url

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    def extract(self) -> ExtractorResult:
        content, img_list = self.parse_docx(self.context.file_bytes)
        return ExtractorResult(
            md_content=content,
            documents=[Document(page_content=content, metadata={"source": self.context.file_name})],
            img_list=img_list,
        )

    def _extract_images_from_docx(self, doc):
        image_map = {}
        img_list = []
        image_service = self.context.image_service
        if image_service is None:
            return image_map, img_list

        for relationship_id, relationship in doc.part.rels.items():
            if "image" not in relationship.target_ref:
                continue
            if relationship.is_external:
                asset = image_service.download_and_upload(relationship.target_ref)
                map_key = relationship_id
            else:
                extension = relationship.target_ref.rsplit(".", 1)[-1]
                asset = image_service.upload_embedded(
                    relationship.target_part.blob,
                    extension=extension,
                    source=f"{self.context.file_name}:{relationship.target_ref}",
                )
                map_key = relationship.target_part
            if asset:
                image_map[map_key] = asset.markdown
                img_list.append(asset.file)
            elif relationship.is_external and is_http_url(relationship.target_ref):
                image_map[map_key] = f"![image]({relationship.target_ref})"

        return image_map, img_list

    def _table_to_markdown(self, table, image_map):
        if not table.rows:
            return ""
        total_cols = max(len(row.cells) for row in table.rows)
        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        rows = [self._parse_row(row, image_map, total_cols) for row in table.rows[1:]]
        return render_markdown_table(headers, rows)

    def _parse_row(self, row, image_map, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            # make sure the col_index is not out of range
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not image_id:
                        continue

                    if image_id in paragraph.part.rels:
                        rel = paragraph.part.rels[image_id]
                        if rel.is_external:
                            if image_id in image_map:
                                paragraph_content.append(image_map[image_id])
                        else:
                            image_part = rel.target_part
                            if image_part in image_map:
                                image_link = image_map[image_part]
                                paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def parse_docx(self, file_bytes):
        doc = DocxDocument(BytesIO(file_bytes))

        content = []

        image_map, img_list = self._extract_images_from_docx(doc)

        def parse_paragraph(paragraph):
            def append_image_link(image_id, has_drawing, target_buffer):
                """Helper to append image link from image_map based on relationship type."""
                rel = doc.part.rels[image_id]
                if rel.is_external:
                    if image_id in image_map and not has_drawing:
                        target_buffer.append(image_map[image_id])
                else:
                    image_part = rel.target_part
                    if image_part in image_map and not has_drawing:
                        target_buffer.append(image_map[image_part])

            def process_run(run, target_buffer):
                # Helper to extract text and embedded images from a run element and append them to target_buffer
                if (
                    hasattr(run.element, "tag")
                    and isinstance(run.element.tag, str)
                    and run.element.tag.endswith("r")
                ):
                    # Process drawing type images
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    has_drawing = False
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                rel = doc.part.rels.get(embed_id)
                                if rel is not None and rel.is_external:
                                    # External image: use embed_id as key
                                    if embed_id in image_map:
                                        has_drawing = True
                                        target_buffer.append(image_map[embed_id])
                                else:
                                    # Internal image: use target_part as key
                                    image_part = doc.part.related_parts.get(embed_id)
                                    if image_part in image_map:
                                        has_drawing = True
                                        target_buffer.append(image_map[image_part])
                    # Process pict type images
                    shape_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
                    )
                    for shape in shape_elements:
                        # Find image data in VML
                        shape_image = shape.find(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}binData"
                        )
                        if shape_image is not None and shape_image.text:
                            image_id = shape_image.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                            )
                            if image_id and image_id in doc.part.rels:
                                append_image_link(image_id, has_drawing, target_buffer)
                        # Find imagedata element in VML
                        image_data = shape.find(".//{urn:schemas-microsoft-com:vml}imagedata")
                        if image_data is not None:
                            image_id = image_data.get("id") or image_data.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                            )
                            if image_id and image_id in doc.part.rels:
                                append_image_link(image_id, has_drawing, target_buffer)
                if run.text:
                    target_buffer.append(run.text)

            def process_hyperlink(hyperlink_elem, target_buffer):
                # Helper to extract text from a hyperlink element and append it to target_buffer
                r_id = hyperlink_elem.get(qn("r:id"))

                # Extract text from runs inside the hyperlink
                link_text_parts = []
                for run_elem in hyperlink_elem.findall(qn("w:r")):
                    run = Run(run_elem, paragraph)
                    # Hyperlink text may be split across multiple runs (e.g., with different formatting),
                    # so collect all run texts first
                    if run.text:
                        link_text_parts.append(run.text)

                link_text = "".join(link_text_parts).strip()

                # Resolve URL
                if r_id:
                    try:
                        rel = doc.part.rels.get(r_id)
                        if rel and rel.is_external:
                            link_text = f"[{link_text or rel.target_ref}]({rel.target_ref})"
                    except Exception:
                        logger.exception("Failed to resolve URL for hyperlink with r:id: %s", r_id)

                if link_text:
                    target_buffer.append(link_text)

            paragraph_content = []
            # State for legacy HYPERLINK fields
            hyperlink_field_url = None
            hyperlink_field_text_parts: list = []
            is_collecting_field_text = False
            # Iterate through paragraph elements in document order
            for child in paragraph._element:
                tag = child.tag
                if tag == qn("w:r"):
                    # Regular run
                    run = Run(child, paragraph)

                    # Check for fldChar (begin/end/separate) and instrText for legacy hyperlinks
                    fld_chars = child.findall(qn("w:fldChar"))
                    instr_texts = child.findall(qn("w:instrText"))

                    # Handle Fields
                    if fld_chars or instr_texts:
                        # Process instrText to find HYPERLINK "url"
                        for instr in instr_texts:
                            if instr.text and "HYPERLINK" in instr.text:
                                # Quick regex to extract URL
                                match = re.search(
                                    r'HYPERLINK\s+"([^"]+)"', instr.text, re.IGNORECASE
                                )
                                if match:
                                    hyperlink_field_url = match.group(1)

                        # Process fldChar
                        for fld_char in fld_chars:
                            fld_char_type = fld_char.get(qn("w:fldCharType"))
                            if fld_char_type == "begin":
                                # Start of a field: reset legacy link state
                                hyperlink_field_url = None
                                hyperlink_field_text_parts = []
                                is_collecting_field_text = False
                            elif fld_char_type == "separate":
                                # Separator: if we found a URL, start collecting visible text
                                if hyperlink_field_url:
                                    is_collecting_field_text = True
                            elif fld_char_type == "end":
                                # End of field
                                if is_collecting_field_text and hyperlink_field_url:
                                    # Create markdown link and append to main content
                                    display_text = "".join(hyperlink_field_text_parts).strip()
                                    if display_text:
                                        link_md = f"[{display_text}]({hyperlink_field_url})"
                                        paragraph_content.append(link_md)
                                # Reset state
                                hyperlink_field_url = None
                                hyperlink_field_text_parts = []
                                is_collecting_field_text = False

                    # Decide where to append content
                    target_buffer = (
                        hyperlink_field_text_parts
                        if is_collecting_field_text
                        else paragraph_content
                    )
                    process_run(run, target_buffer)
                elif tag == qn("w:hyperlink"):
                    process_hyperlink(child, paragraph_content)
            return "".join(paragraph_content).strip() if paragraph_content else ""

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph.strip():
                        content.append(parsed_paragraph)
                    else:
                        content.append("\n")
                elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # table
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table, image_map))
        return "\n".join(content).strip(), img_list
